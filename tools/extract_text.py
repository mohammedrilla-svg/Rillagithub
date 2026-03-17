import os


def extract_text_from_file(file_path):
    """
    Extracts text content from PDF, Word (.docx), or Excel (.xlsx) files.
    Returns: dict with 'text' or 'error'
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.pdf':
            return extract_pdf(file_path)
        elif ext == '.docx':
            return extract_docx(file_path)
        elif ext in ('.xlsx', '.xls'):
            return extract_excel(file_path)
        else:
            return {"error": f"Unsupported file type: {ext}. Supported: .pdf, .docx, .xlsx"}
    except Exception as e:
        return {"error": f"Failed to extract text: {str(e)}"}


def extract_pdf(file_path):
    import pdfplumber
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    if not text.strip():
        return {"error": "Could not extract any text from PDF. It may be image-based."}
    return {"text": text.strip()}


def extract_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += "\n" + row_text
    if not text.strip():
        return {"error": "Could not extract any text from Word document."}
    return {"text": text.strip()}


def extract_excel(file_path):
    from openpyxl import load_workbook
    wb = load_workbook(file_path, read_only=True)
    text = ""
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text += f"Sheet: {sheet_name}\n"
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text:
                text += row_text + "\n"
        text += "\n"
    wb.close()
    if not text.strip():
        return {"error": "Could not extract any text from Excel file."}
    return {"text": text.strip()}
